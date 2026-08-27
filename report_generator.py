#!/usr/bin/env python3
"""
Report generator for Bentley Compass 360.
Generates Word documents for Self-Assessment, Full 360, and Progress Reports.

Updated for 9 dimensions (45 items total) with Performance Excellence dimension.
The old scored Overall Effectiveness items (Q46-47) are replaced by two open
text prompts ("keep" / "change"), folded into Overall Qualitative Feedback.
"""

import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')
from datetime import datetime
import tempfile
import os
import uuid
import sys
import time
import textwrap
import requests
from pathlib import Path

from framework import (
    DIMENSIONS, DIMENSION_DESCRIPTIONS,
    COLOURS, GROUP_COLOURS, GROUP_DISPLAY,
    HIGH_SCORE_THRESHOLD, SIGNIFICANT_GAP,
    REPORT_FONT, CHART_FONT_PREFERENCE,
    BENTLEY_CHART_REGULAR_PATH, BENTLEY_CHART_LIGHT_PATH,
    BENTLEY_CHART_EXPANDED_BOLD_PATH,
    LOGO_PATH,
    get_item_text
)

REPORTS_DIR = Path("reports")
REPORTS_DIR.mkdir(exist_ok=True)

# LOGO_PATH itself lives in framework.py - the single source of truth for the
# asset path, shared with the Streamlit pages and email templates. Report
# generation degrades gracefully (skips the picture, prints a one-line stderr
# warning) if the file is ever missing, rather than crashing.
_logo_missing_warned = False


def _add_logo_if_present(paragraph, width_in):
    """Add the Bentley wings logo to a paragraph at width_in inches.

    Caller sets the paragraph's alignment beforehand. Returns True if the
    picture was added. Missing-asset case is reported once per process to
    stderr rather than silently leaving a gap in every report.
    """
    global _logo_missing_warned
    if not LOGO_PATH.exists():
        if not _logo_missing_warned:
            print(
                f"WARNING: logo asset not found at {LOGO_PATH} - reports will "
                f"generate without it until the file is added.",
                file=sys.stderr,
            )
            _logo_missing_warned = True
        return False
    run = paragraph.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(width_in))
    return True

# ============================================
# PAGE GEOMETRY
# ============================================
#
# A4, set on the human's instruction 2026-08-04. python-docx's default
# template is US LETTER, which would shift the margins when UK readers print
# it, so `apply_page_geometry` overrides it.
#
# Margins reduced 2026-08-07 from a uniform 1in - checked first whether that
# figure was a tested/specified constraint (print testing, a known printer
# compatibility issue, a client spec) and found none: it was set alongside
# the A4 fix above as a conventional round default, not for any margin-
# specific reason. 0.5in top/bottom and 0.6in left/right is still a
# conservative, standard professional-document margin, comfortably within
# safe printing tolerances on any normal home or office printer - nothing
# like the sub-0.1in margins seen in the reference document that prompted
# this. Left/right slightly wider than top/bottom to leave room for
# hole-punching or binding if these ever get printed and filed.
#
# Derive everything from these constants rather than hard-coding inch values, so
# a future page change does not silently leave tables the wrong width.
PAGE_WIDTH_IN = 8.27      # A4 = 210mm
PAGE_HEIGHT_IN = 11.69    # A4 = 297mm
MARGIN_TOP_BOTTOM_IN = 0.5
MARGIN_LEFT_RIGHT_IN = 0.6

CONTENT_WIDTH_IN = PAGE_WIDTH_IN - (2 * MARGIN_LEFT_RIGHT_IN)    # 7.07in
CONTENT_HEIGHT_IN = PAGE_HEIGHT_IN - (2 * MARGIN_TOP_BOTTOM_IN)  # 10.69in

# Item bar chart width - FIXED, deliberately NOT derived from CONTENT_WIDTH_IN.
# It was proportional to content width until 2026-08-07, which seemed harmless
# (charts filling available space is the right instinct for most elements),
# but bar charts are raster images with a fixed aspect ratio: widening one
# widens it proportionally taller too. The 2026-08-07 margin reduction grew
# CONTENT_WIDTH_IN and, as a side effect, grew every chart's rendered height
# by the same ~14% - enough on its own to push a 5-item dimension down to 4
# per page, with zero change to any spacing value. Fixing the width breaks
# that coupling: the item text column absorbs whatever extra horizontal room
# a page-geometry change provides instead, and chart height stays predictable
# regardless of margins. 2.9in is the value verified (see
# _estimate_five_items_fit below and the 2026-08-07 render check) to keep a
# worst-case 6-bar item's chart short enough that five of them, the heading,
# and the description all fit within CONTENT_HEIGHT_IN at current spacing.
ITEM_CHART_WIDTH_IN = 2.9

# Model used for the AI theme synthesis section.
#
# WAS `claude-sonnet-4-20250514`, which is DEPRECATED with a retirement date of
# 2026-06-15 — already past. A retired model ID returns 404, and the synthesis
# function swallows that into a silent skip, so the Key Themes section would have
# quietly vanished from every report even with a valid API key configured.
#
# Two behaviours of the current model that the request body has to account for:
#   - Thinking is ON BY DEFAULT, and max_tokens caps thinking plus response text
#     together, so max_tokens needs real headroom (see the call site).
#   - Sampling parameters (temperature/top_p/top_k) are rejected with a 400.
#     Do not add them.
# Swap to 'claude-sonnet-5' if the human decides the cost per report matters more
# than synthesis quality.
SYNTHESIS_MODEL = 'claude-opus-5'

# HTTP statuses worth retrying: 429 (rate limited), 500 (api_error), 529
# (overloaded_error) are all transient, server-side, and routinely resolve a
# few seconds later. 4xx statuses like 400/401/403/404 mean something is wrong
# with the key or request and retrying won't help, so they fail immediately.
SYNTHESIS_TRANSIENT_STATUS_CODES = {429, 500, 529}
SYNTHESIS_MAX_RETRIES = 2
SYNTHESIS_RETRY_DELAY_SECONDS = 3


def content_columns(*relative_widths):
    """
    Column widths that fill the content width exactly, from relative proportions.

    Using proportions rather than literal inches means the tables stay aligned
    with each other and with the charts if the page size ever changes again.
    """
    total = sum(relative_widths)
    return [Inches(CONTENT_WIDTH_IN * w / total) for w in relative_widths]


def apply_page_geometry(doc):
    """Set A4 page size and margins on every section of the document."""
    for section in doc.sections:
        section.page_width = Inches(PAGE_WIDTH_IN)
        section.page_height = Inches(PAGE_HEIGHT_IN)
        section.left_margin = Inches(MARGIN_LEFT_RIGHT_IN)
        section.right_margin = Inches(MARGIN_LEFT_RIGHT_IN)
        section.top_margin = Inches(MARGIN_TOP_BOTTOM_IN)
        section.bottom_margin = Inches(MARGIN_TOP_BOTTOM_IN)


# ============================================
# TYPOGRAPHY
# ============================================

from matplotlib import font_manager as _fm


def _load_chart_font(path):
    """FontProperties for one Bentley chart-font file, loaded BY PATH, not
    by name. See BENTLEY_CHART_*_PATH in framework.py for why: these three
    "_web.ttf" files have a correct family name on their Macintosh
    name-table entry, but the Windows-platform entries matplotlib/freetype
    actually reads for name-based lookup are scrubbed to the literal DEL
    character - a real, verified (not assumed) name-table quirk, not a
    corrupted or substitute font. FontProperties(fname=...) bypasses name
    resolution entirely and loads the file's own glyph data directly, so
    the corruption doesn't matter. Returns None if the file is missing, so
    callers degrade to matplotlib's default rather than crash on a missing
    asset - same pattern as get_logo_data_uri/font_face_css elsewhere.
    """
    if not path.exists():
        return None
    try:
        return _fm.FontProperties(fname=str(path))
    except Exception:
        return None


# Registered once at import time, reused across every chart this module
# generates - loading a font file repeatedly per-chart would be wasteful
# and gains nothing, the file's contents don't change between calls.
CHART_FONT_REGULAR = _load_chart_font(BENTLEY_CHART_REGULAR_PATH)
CHART_FONT_LIGHT = _load_chart_font(BENTLEY_CHART_LIGHT_PATH)
CHART_FONT_BOLD = _load_chart_font(BENTLEY_CHART_EXPANDED_BOLD_PATH)


def chart_font(size, bold=False):
    """A correctly-SIZED copy of the shared Bentley chart FontProperties,
    for one text call. REQUIRED, not a convenience: confirmed live that
    matplotlib's Text/tick-label API applies `fontproperties` AFTER any
    separate `size=`/`fontsize=` kwarg, so passing both `size=14` and
    `fontproperties=CHART_FONT_BOLD` to the same call silently DISCARDS the
    size and falls back to the FontProperties object's own default (10pt) -
    confirmed by actually rendering and reading back label.get_size(), not
    assumed from the matplotlib docs. The size has to be baked into the
    FontProperties object itself instead. Always returns a COPY
    (FontProperties.copy(), not the shared module-level singleton) -
    mutating CHART_FONT_BOLD's size in place would leak into every other
    call site that also uses it. Falls back to a plain size dict (which
    matplotlib also accepts wherever a FontProperties is expected) if the
    underlying font file was missing at import time, so callers don't need
    a separate None-check.
    """
    base = CHART_FONT_BOLD if bold else CHART_FONT_REGULAR
    if base is None:
        return {'size': size}
    prop = base.copy()
    prop.set_size(size)
    return prop

# Every text call in create_radar_chart/create_item_bar_chart/
# create_self_only_bar passes fontproperties=CHART_FONT_REGULAR or
# CHART_FONT_BOLD explicitly now, matching whichever weight that element
# already used (fontweight='bold' -> CHART_FONT_BOLD, otherwise ->
# CHART_FONT_REGULAR) - never BOTH fontproperties and fontweight together,
# since asking matplotlib to also synthesise bold on top of an
# already-bold custom face is not what's wanted and was never tested.
# CHART_FONT_LIGHT is registered for completeness (Light is one of the two
# weights the UI typeface work already established) but nothing in these
# three chart functions currently has a moment that calls for it - none of
# their existing text was ever set to a lighter-than-default weight.


def resolve_chart_font():
    """
    LEGACY name-based fallback - only reached if BENTLEY_CHART_REGULAR_PATH
    is missing from disk entirely (the primary CHART_FONT_REGULAR above is
    then None), so there's still a chance of finding a real Bentley face
    if one happens to be installed system-wide the old way (true on Ian's
    own Mac - see the CLAUDE.md environment-gotchas entry this used to be
    the whole story for). Not the primary mechanism any more: fname-based
    loading above works identically on every machine regardless of what's
    installed there, which is the actual point of this rework - closing
    the local-vs-Render generation gap, not just working around it.
    """
    try:
        available = {f.name for f in _fm.fontManager.ttflist}
    except Exception:
        return None

    for candidate in CHART_FONT_PREFERENCE:
        if candidate in available:
            return candidate
    return None


if CHART_FONT_REGULAR is None:
    _name_fallback_font = resolve_chart_font()
    if _name_fallback_font:
        plt.rcParams['font.family'] = _name_fallback_font


def apply_document_font(doc, font_name=REPORT_FONT):
    """
    Set the font on the document's base and heading styles.

    Runs throughout this module set size and colour but not the typeface, so they
    inherit from these styles. Word needs the East Asian attribute set as well,
    otherwise it silently keeps its own default for some runs.
    """
    style_names = [
        'Normal', 'Title',
        'Heading 1', 'Heading 2', 'Heading 3',
        'List Bullet', 'List Number',
    ]
    # Heading 1-3 inherit Word's default template colour (a theme blue) unless
    # given an explicit override — setting the font name/size above doesn't
    # touch colour. Every section heading in this module goes through
    # doc.add_heading at level 1 or 2, so both need it; level 3 and Title
    # aren't currently used but are included for consistency.
    heading_styles = {'Title', 'Heading 1', 'Heading 2', 'Heading 3'}

    for style_name in style_names:
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue

        style.font.name = font_name
        if style_name in heading_styles:
            style.font.color.rgb = RGBColor(0x18, 0x33, 0x19)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is None:
            rfonts = OxmlElement('w:rFonts')
            rpr.append(rfonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rfonts.set(qn(attr), font_name)

# Fixed display order for grouped comments (2026-08-06), matching the human's
# explicit sequence: Line Manager, Direct Reports, Peers, Others. Self isn't
# part of that sequence but is a real possible source here too - a
# per-dimension comment recorded at self-assessment persists into the Full
# 360's data - so it's placed first rather than silently dropped if present.
COMMENT_GROUP_ORDER = ['Self', 'Boss', 'DRs', 'Peers', 'Others']

# Colour map for comment source labels (RGB tuples for python-docx), matching
# the fixed report palette in framework.COLOURS.
COMMENT_SOURCE_COLOURS = {
    'Line Manager':   RGBColor(0x3D, 0x3D, 0x3D),   # grey_dark
    'Peers':          RGBColor(0x6B, 0x6B, 0x6B),   # grey_mid
    'Direct Reports': RGBColor(0x9A, 0x9A, 0x9A),   # grey_light
    'Others':         RGBColor(0xC4, 0xC4, 0xC4),   # grey_lightest
    'Self':           RGBColor(0x18, 0x33, 0x19),   # Bentley green
}


def set_cell_shading(cell, color):
    """Set background colour of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_section_heading(doc, text, font_size=18, level=1):
    """
    Add a heading with an explicit font size, rather than leaving it to
    inherit whatever Word's built-in Heading-N style default happens to be.

    Dimension names and the PAPU-NANU quadrant headings (Agreed Strengths,
    Good News, Development Areas, Potential Blind Spots) go through this at
    level=2 as of 2026-08-06 - previously raw doc.add_heading(..., level=2)
    calls with no size override, which rendered at 13pt only because that's
    Word's template default for Heading 2, not because anyone chose 13pt.
    """
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(font_size)
    return heading


def make_table_borderless(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)


# ============================================
# CLEAN COMMENT FORMATTING
# ============================================

def _add_thin_rule(doc, colour='CCCCCC', space_pt=2):
    """Add a thin horizontal rule (bottom border on an empty paragraph)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_pt)
    p.paragraph_format.space_after = Pt(space_pt)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{colour}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


# ============================================
# WRITING SPACE
# These make the report a live working document rather than a read-only output:
# the leader annotates it during and after their coaching conversation. Ruled
# lines work whether the report is printed and handwritten on, or typed into.
# ============================================

def add_writing_lines(doc, count=3, indent=Inches(0.25)):
    """Add blank ruled lines for the reader to write or type on."""
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = indent
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)


def add_writing_prompt(doc, text):
    """Add an italic instruction introducing a writing space."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.keep_with_next = True
    run = para.add_run(text)
    run.font.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)
    return para


def add_priority_capture_table(doc, rows=3):
    """
    Add an empty Dimension / actions table for priorities the leader ADDS during
    their coaching conversation.

    Deliberately mirrors the layout of the stated priorities above it, so the
    additions read as part of the same list rather than a separate exercise.
    These are captured in the document, not in the database: the stored
    priorities are the pre-feedback set, and preserving those unchanged is what
    makes the comparison in the Full 360 report meaningful.
    """
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False

    widths = content_columns(2.0, 4.0)

    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "What I will do differently"
    for i, cell in enumerate(hdr):
        cell.width = widths[i]
        set_cell_shading(cell, '183319')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)

    for _ in range(rows):
        row = table.add_row().cells
        for i, cell in enumerate(row):
            cell.width = widths[i]
            # Give each empty row some height to write in
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(9)
            para.paragraph_format.space_after = Pt(9)

    return table


def _add_comment_group_block(doc, group_code, comment_texts):
    """
    Add one rater-type heading followed by every comment from that group -
    the heading appears once per group, not once per comment. Individual
    comments within the group are separated by a lighter divider than the
    one add_clean_comments uses between groups, so the hierarchy (group >
    individual comment) is visually clear rather than reading as a flat list.
    """
    display_name = GROUP_DISPLAY.get(group_code, group_code)

    # Source label - keep_with_next stops Word breaking the page right here,
    # which would otherwise strand the label alone at the bottom of a page
    # with its first comment starting fresh on the next.
    source_para = doc.add_paragraph()
    source_para.paragraph_format.space_before = Pt(8)
    source_para.paragraph_format.space_after = Pt(2)
    source_para.paragraph_format.keep_with_next = True
    run = source_para.add_run(display_name)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = COMMENT_SOURCE_COLOURS.get(
        display_name, RGBColor(0x5F, 0x5E, 0x5A)
    )

    for i, comment_text in enumerate(comment_texts):
        comment_para = doc.add_paragraph()
        comment_para.paragraph_format.space_before = Pt(0)
        # Tighter than the previous 6pt: this paragraph's own trailing space
        # only needs to clear the (now equally tight) divider that follows,
        # not carry the whole visual gap on its own.
        comment_para.paragraph_format.space_after = Pt(3)
        comment_para.paragraph_format.line_spacing = 1.0
        run = comment_para.add_run(comment_text)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

        if i < len(comment_texts) - 1:
            # Lighter than the CCCCCC rule between groups, and tighter
            # (space_pt=1 not the default 2) - a quieter, closer divider for
            # comments that share the same source, versus the more clearly
            # separated one between different rater-type groups.
            _add_thin_rule(doc, colour='E8E8E8', space_pt=1)


def add_clean_comments(doc, comments_list):
    """
    Add a set of comments grouped by rater type: one heading per group, in
    the fixed order COMMENT_GROUP_ORDER, with every comment from that group
    listed beneath it. comments_list: list of dicts with 'group' and 'text'
    keys - previously rendered one heading per comment, which repeated the
    same rater-type heading back to back whenever more than one person in a
    group left a comment.
    """
    if not comments_list:
        return

    grouped = {}
    for comment in comments_list:
        grouped.setdefault(comment['group'], []).append(comment['text'])

    # keep_with_next on the opening rule chains it to the first group's own
    # label (which in turn keeps itself with its first comment - see
    # _add_comment_group_block), so a page break can't land between the
    # "Comments on X" heading and the first comment either.
    opening_rule = _add_thin_rule(doc)
    opening_rule.paragraph_format.keep_with_next = True

    is_first_group = True
    for group_code in COMMENT_GROUP_ORDER:
        comment_texts = grouped.get(group_code)
        if not comment_texts:
            continue
        if not is_first_group:
            _add_thin_rule(doc)  # heavier divider between groups
        _add_comment_group_block(doc, group_code, comment_texts)
        is_first_group = False

    _add_thin_rule(doc)


# ============================================
# PAPU-NANU CATEGORISATION
# ============================================

def categorize_papu_nanu(data):
    """Categorise items into PAPU-NANU quadrants."""
    categories = {
        'agreed_strengths': [],
        'good_news': [],
        'development_areas': [],
        'hidden_talents': [],
    }
    
    for item_num, item_scores in data['by_item'].items():
        self_score = item_scores.get('Self')
        combined = item_scores.get('Combined')
        gap = item_scores.get('Gap')
        
        if self_score is None or combined is None:
            continue
        
        no_opp_info = data.get('no_opportunity', {}).get(item_num)
        no_opp_count = no_opp_info['count'] if no_opp_info else 0
        
        item_info = {
            'item_num': item_num,
            'text': item_scores.get('text', get_item_text(item_num, 'Others')),
            'self': self_score,
            'combined': combined,
            'gap': gap,
            'no_opp_count': no_opp_count,
        }
        
        if combined >= HIGH_SCORE_THRESHOLD:
            if gap is not None and gap < -SIGNIFICANT_GAP:
                categories['good_news'].append(item_info)
            elif gap is not None and gap > SIGNIFICANT_GAP:
                categories['hidden_talents'].append(item_info)
            else:
                categories['agreed_strengths'].append(item_info)
        else:
            if gap is not None and gap > SIGNIFICANT_GAP:
                categories['hidden_talents'].append(item_info)
            else:
                categories['development_areas'].append(item_info)
    
    for cat in categories:
        if cat in ['agreed_strengths', 'good_news']:
            categories[cat].sort(key=lambda x: x['combined'], reverse=True)
        else:
            categories[cat].sort(key=lambda x: x['combined'])
    
    return categories


# ============================================
# CHARTS
# ============================================

def create_radar_chart(dimensions, self_scores, combined_scores, output_path):
    """Create radar chart for dimension overview - professional style."""
    labels = list(dimensions.keys())
    num_vars = len(labels)
    
    # Calculate angles - start at top and go CLOCKWISE
    # We need to go from 90 degrees (top) decreasing (clockwise)
    angles_deg = [90 - (360 * i / num_vars) for i in range(num_vars)]
    angles = [np.radians(a) for a in angles_deg]
    angles += angles[:1]  # Complete the circle
    
    # Get values
    self_values = [self_scores.get(dim, 0) or 0 for dim in labels]
    self_values += self_values[:1]
    
    # Create figure
    fig, ax = plt.subplots(figsize=(12, 12), subplot_kw=dict(polar=True))
    
    # Set theta to start from top and go clockwise
    ax.set_theta_offset(np.pi / 2)  # Start from top
    ax.set_theta_direction(-1)  # Go clockwise
    
    # Recalculate angles for clockwise from top (simpler now with direction set)
    angles = np.linspace(0, 2 * np.pi, num_vars, endpoint=False).tolist()
    angles += angles[:1]
    
    # Get values again with new angle order
    self_values = [self_scores.get(dim, 0) or 0 for dim in labels]
    self_values += self_values[:1]
    
    # Style the grid
    ax.set_facecolor('white')
    ax.spines['polar'].set_color('#999999')
    ax.spines['polar'].set_linewidth(1.5)
    ax.grid(color='#999999', linestyle='-', linewidth=1, alpha=0.8)
    
    # Plot Self scores
    ax.plot(angles, self_values, 'o-', linewidth=3, label='Self', 
            color=COLOURS['bentley_green'], markersize=10)
    ax.fill(angles, self_values, alpha=0.25, color=COLOURS['bentley_green'])
    
    # Plot Combined scores if available. FILL uses 'heritage_white' - matches
    # the "All Raters" bar chart colour exactly, and is what actually
    # delivers the visual identity match, since the large shaded region is
    # what a reader associates with a colour. Composited-luminance maths
    # confirms it doesn't reintroduce the overlap-muddiness problem this
    # chart's two stacked semi-transparent fills had before (Self-only vs
    # All-Raters-only luminance gap is 0.168 with heritage_white, wider than
    # the old grey fill's 0.132). LINE/markers use 'heritage_white_deep', not
    # the base heritage_white: a full-opacity heritage_white line measures
    # only 1.44:1 contrast against the white chart background, effectively
    # invisible, and even the deepened variant only reaches 2.0:1 (short of
    # the ~3:1 WCAG guideline for graphical elements) - but matching the All
    # Raters bar identity matters more here than maximising line contrast,
    # per the human's explicit call 2026-08-08, given the fill above already
    # carries genuine heritage_white and is what most visibly makes the
    # identity match.
    if combined_scores and any(combined_scores.get(dim) for dim in labels):
        combined_values = [combined_scores.get(dim, 0) or 0 for dim in labels]
        combined_values += combined_values[:1]
        ax.plot(angles, combined_values, 'o-', linewidth=3, label='All Raters',
                color=COLOURS['heritage_white_deep'], markersize=10)
        ax.fill(angles, combined_values, alpha=0.25, color=COLOURS['heritage_white'])
    
    # Configure the chart
    ax.set_ylim(0, 5)
    ax.set_yticks([1, 2, 3, 4, 5])
    ax.set_yticklabels(['1', '2', '3', '4', '5'], color='#333333',
                        fontproperties=chart_font(14, bold=True))
    # Sit the 1-5 scale in the gap between the first two spokes rather than on
    # top of a dimension label. With theta offset pi/2 and clockwise direction,
    # a data angle of d appears at screen angle (90 - d), so 20 lands midway
    # between the first spoke (top) and the second.
    ax.set_rlabel_position(20)

    # Wrap long dimension names. Without this, a label like "Building
    # High-Performing Teams" is centred on its spoke and its horizontal extent
    # runs back across the plot area.
    wrapped = ['\n'.join(textwrap.wrap(label, width=16)) for label in labels]

    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(wrapped, color='#333333',
                        fontproperties=chart_font(13, bold=True))
    ax.tick_params(axis='x', pad=18)

    # Anchor each label on the side facing away from the chart, so text grows
    # outwards instead of inwards. Labels on the right half are left-aligned,
    # the left half right-aligned, and the top and bottom stay centred.
    for label, angle in zip(ax.get_xticklabels(), angles[:-1]):
        screen_deg = (90 - np.degrees(angle)) % 360
        if np.isclose(screen_deg, 90) or np.isclose(screen_deg, 270):
            label.set_horizontalalignment('center')
        elif screen_deg < 90 or screen_deg > 270:
            label.set_horizontalalignment('left')
        else:
            label.set_horizontalalignment('right')

    # Add legend
    #
    # bold=True ADDED 2026-08-27 - Ian caught this by eye: with Regular
    # weight, "Self"/"All Raters" reads visibly plainer than every other
    # label on this same chart (spoke labels and the 1-5 scale are all
    # Expanded Bold), and Regular alone doesn't carry much of Bentley's
    # visual distinctiveness at a glance - confirmed the font WAS
    # genuinely being applied even before this change (pixel-diffed the
    # rendered "All Raters" text against matplotlib's own default font at
    # the same size - 10,624 differing pixels, not a silent fallback),
    # it just didn't look distinctively branded next to the bold text
    # around it. Matching the rest of the chart's weight, not fixing a
    # bug.
    if combined_scores and any(combined_scores.get(dim) for dim in labels):
        ax.legend(loc='lower center', bbox_to_anchor=(0.5, -0.13),
                  ncol=2, prop=chart_font(14, bold=True), frameon=False)

    # Leave room around the plot for the labels. tight_layout fights with polar
    # axes that have long outward labels, so set the margins explicitly.
    fig.subplots_adjust(left=0.18, right=0.82, top=0.86, bottom=0.14)
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white', pad_inches=0.35)
    plt.close()


def create_item_bar_chart(scores, output_path, include_combined=True):
    """Create horizontal bar chart for an item showing all respondent groups."""
    groups = []
    values = []
    colors = []
    
    for group in ['Self', 'Boss', 'Peers', 'DRs', 'Others']:
        val = scores.get(group)
        if val is not None:
            groups.append(GROUP_DISPLAY[group])
            values.append(val)
            colors.append(GROUP_COLOURS[group])
    
    # Add combined bar if requested and available. Heritage White, matching
    # the radar chart's All Raters fill (see create_radar_chart) - both now
    # use the same colour identity.
    if include_combined and scores.get('Combined') is not None:
        groups.append(GROUP_DISPLAY['Combined'])
        values.append(scores['Combined'])
        colors.append(COLOURS['heritage_white'])
    
    if not values:
        return False
    
    groups = groups[::-1]
    values = values[::-1]
    colors = colors[::-1]
    
    # Per-group vertical allocation tightened again from 0.4in to 0.32in.
    # The "5 items always fit on one page" guarantee (see add_dimension_section
    # and the w:cantSplit on each item row) was only ever verified against data
    # where some groups had folded into Others, i.e. 3-4 bars per item. A real,
    # fully unfolded item (Self + Boss + Peers + DRs + Others + Combined, 6 bars)
    # is taller, and on the FIRST dimension's page - which also carries the
    # "Detailed Feedback by Dimension" heading the other dimensions don't have -
    # that extra height was enough to push the 5th item onto a page of its own.
    # Confirmed on a real 6-bar report: Leading Self fit only 4 of 5 items before
    # this change.
    #
    # Floor raised from 0.65in to 0.8in when doing this - NOT left alone. The
    # floor only ever bound the 1-bar case under the old 0.4 multiplier
    # (1*0.4=0.4in, floored to 0.65in); a 2-bar item sat at 0.8in, safely above
    # the floor, so the floor's own value was never actually exercised at 2
    # bars. Dropping the multiplier to 0.32 put 2-bar items AT the old 0.65in
    # floor for the first time (2*0.32=0.64in) - and 0.65in turned out to be too
    # short for 2 bars plus their value labels: rendered and found the two
    # labels overlapping. Raising the floor to 0.8in restores the 2-bar case to
    # its previous, working height exactly; a 1-bar item now gets slightly more
    # room than before (0.65in -> 0.8in), which only ever makes that sparser
    # case easier to read, not harder.
    #
    # CORRECTION: shrinking the slot from 0.4in to 0.32in without touching
    # `height` below cut the bar itself, not just the gap - at height=0.4 the
    # bar is a fixed 40% of whatever the slot is (0.16in old, 0.128in new),
    # so both bar and gap shrank by the same 20%. The actual instruction was
    # to leave bar depth alone and take the whole cut out of the whitespace
    # between bars. Raised height 0.4 -> 0.5 to compensate: 0.5 * 0.32in =
    # 0.16in, the SAME bar thickness as before the slot was ever reduced. The
    # gap absorbs all of it instead: 0.6*0.4in=0.24in -> 0.5*0.32in=0.16in,
    # a 33% cut to whitespace only. Total slot height (and therefore the page-
    # fit fix) is unchanged by this - height only redistributes space within
    # the slot, it doesn't add to it.
    fig, ax = plt.subplots(figsize=(4.5, max(0.8, len(groups) * 0.32)))

    y_pos = np.arange(len(groups))
    bars = ax.barh(y_pos, values, color=colors, height=0.5)
    
    ax.set_yticks(y_pos)
    ax.set_yticklabels(groups, fontproperties=chart_font(10, bold=True))
    ax.set_xlim(0, 6.0)
    ax.set_xticks([1, 2, 3, 4, 5])
    ax.tick_params(axis='x', labelsize=10)
    # tick_params only controls size/colour, not font family - the x-axis's
    # numeric labels (1-5) are auto-generated by set_xticks rather than an
    # explicit set_xticklabels call, so they need setting directly here or
    # they're left on matplotlib's default font.
    for tick_label in ax.get_xticklabels():
        tick_label.set_fontproperties(chart_font(10))

    ax.axvline(x=4, color='green', linestyle='--', alpha=0.3, linewidth=1)
    ax.axvline(x=3, color='gray', linestyle=':', alpha=0.3, linewidth=1)

    # Place all scores at fixed right-aligned position
    for bar, val in zip(bars, values):
        ax.text(5.7, bar.get_y() + bar.get_height()/2, f'{val:.1f}',
                va='center', ha='right', fontproperties=chart_font(12, bold=True))
    
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', pad_inches=0.05, facecolor='white')
    plt.close()
    return True


def create_self_only_bar(score, output_path):
    """Create horizontal bar chart for self-assessment only."""
    # figsize height raised 0.65 -> 0.72 (2026-08-27, Ian's own live catch:
    # "the bar graphs rendered for the Self-Assessment report are thinner
    # than the bars in the Full 360... plenty of page real estate... to
    # thicken them to the same depth"). The earlier 0.4->0.6 bar-height-
    # FRACTION change (comment below) improved things but never actually
    # closed the gap to Full 360, because bar-height is a fraction of the
    # PLOT AREA, and matplotlib reserves a roughly fixed amount of the
    # figure for the x-axis tick labels regardless of total figure height
    # - at a short 0.65in figure, most of that height was tick-label
    # margin, leaving very little for the plot area the bar actually
    # lives in. Confirmed empirically, not guessed: measured actual
    # rendered bar thickness in pixels (searching the saved PNG for
    # bentley_green pixel rows, not eyeballing) - the OLD 0.65in figure
    # produced a 12px-thick bar, against 17px for a REAL densest-case
    # (6-bar) Full 360 item chart's "Self" bar at the same native
    # resolution and the same eventual ITEM_CHART_WIDTH_IN embed scale
    # (both charts share the same 4.5in native width -> 2.9in embed
    # width, so comparing raw pixel thickness between them is a fair,
    # direct comparison of final embedded size). 0.72in was found by
    # sweeping figure heights and re-measuring until the output matched
    # the 6-bar target (18px, 1px over rather than under) - the
    # relationship is NOT linear (0.65->0.85in jumps 12px->30px) because
    # of that fixed tick-label margin, so this was tuned empirically
    # against the real target, not computed analytically. The densest
    # (6-bar, i.e. every group present) Full 360 case was chosen as the
    # matching target deliberately, not the sparser 1-2 bar cases (which
    # render much thicker still, ~39px, due to a separate 0.8in figure-
    # height FLOOR in create_item_bar_chart) - a completed 360 typically
    # shows all groups, so that's the more representative "Full 360 bar"
    # for a leader to visually compare a self-assessment page against.
    # Verified this doesn't cost anything towards the self-assessment's
    # own page-fit after making the change - see the live regeneration
    # note where this was checked, not just assumed from Ian's "plenty
    # of page real estate" alone.
    fig, ax = plt.subplots(figsize=(4.5, 0.72))

    # height raised 0.4 -> 0.6: measured at 0.4 this bar embeds at only
    # ~0.036in (2.6pt) thick once placed in the document at ITEM_CHART_WIDTH_IN
    # - visibly a thin line, and noticeably thinner than the Full 360 item
    # charts' bars (~0.078in/5.6pt in the densest 6-bar case). The figure's
    # total height (0.65in) is untouched, so this costs nothing towards the
    # "5 items per page" layout - height only redistributes the SAME fixed
    # vertical space from margin into bar, it doesn't add any. SUPERSEDED BY
    # THE FIGSIZE CHANGE ABOVE, 2026-08-27 - kept for the historical context
    # on why 0.6 (not the original 0.4) is the bar-height fraction being
    # built on here.
    if score is not None:
        ax.barh([0], [score], color=COLOURS['bentley_green'], height=0.6)
        # Place score at fixed right-aligned position
        ax.text(5.7, 0, f'{score:.1f}', va='center', ha='right',
                fontproperties=chart_font(12, bold=True))

    ax.set_xlim(0, 6.0)
    ax.set_ylim(-0.5, 0.5)
    ax.set_xticks([1, 2, 3, 4, 5])
    ax.tick_params(axis='x', labelsize=10)
    for tick_label in ax.get_xticklabels():
        tick_label.set_fontproperties(chart_font(10))
    ax.set_yticks([])
    
    ax.axvline(x=4, color='green', linestyle='--', alpha=0.3, linewidth=1)
    ax.axvline(x=3, color='gray', linestyle=':', alpha=0.3, linewidth=1)
    
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', pad_inches=0.05, facecolor='white')
    plt.close()


# ============================================
# REPORT SECTIONS
# ============================================

def _add_page_number_footer(section):
    """Add 'Page X of Y' footer to a document section, centre-aligned."""
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Page "
    run = para.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Current page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1 = para.add_run()
    run1._r.append(fldChar1)

    instrText1 = OxmlElement('w:instrText')
    instrText1.set(qn('xml:space'), 'preserve')
    instrText1.text = ' PAGE '
    run2 = para.add_run()
    run2._r.append(instrText1)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3 = para.add_run()
    run3._r.append(fldChar2)

    run4 = para.add_run("1")
    run4.font.size = Pt(8)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5 = para.add_run()
    run5._r.append(fldChar3)

    # " of "
    run6 = para.add_run(" of ")
    run6.font.size = Pt(8)
    run6.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Total pages field
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'begin')
    run7 = para.add_run()
    run7._r.append(fldChar4)

    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = ' NUMPAGES '
    run8 = para.add_run()
    run8._r.append(instrText2)

    fldChar5 = OxmlElement('w:fldChar')
    fldChar5.set(qn('w:fldCharType'), 'separate')
    run9 = para.add_run()
    run9._r.append(fldChar5)

    run10 = para.add_run("1")
    run10.font.size = Pt(8)
    run10.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    fldChar6 = OxmlElement('w:fldChar')
    fldChar6.set(qn('w:fldCharType'), 'end')
    run11 = para.add_run()
    run11._r.append(fldChar6)


def create_cover_page(doc, leader_name, report_type, dealership=None, cohort=None):
    """Create the cover page."""
    for _ in range(2):
        doc.add_paragraph()

    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if _add_logo_if_present(logo_para, width_in=2.3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("BENTLEY COMPASS 360")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(report_type)
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run(leader_name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)
    
    if dealership:
        detail = doc.add_paragraph()
        detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = detail.add_run(dealership)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    if cohort:
        detail = doc.add_paragraph()
        detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = detail.add_run(f"Cohort: {cohort}")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(datetime.now().strftime("%B %Y"))
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    for _ in range(2):
        doc.add_paragraph()

    prog = doc.add_paragraph()
    prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = prog.add_run("Bentley Compass Leadership Programme")
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # Ensure cover page (first section) has NO footer
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = False
    footer = first_section.footer
    footer.is_linked_to_previous = False
    # Clear any default footer content
    for p in footer.paragraphs:
        p.text = ""
    
    # Add a SECTION BREAK (new page) so the rest of the report is a new section
    from docx.enum.section import WD_ORIENT
    new_section = doc.add_section()
    new_section.start_type = 2  # 2 = new page
    # Copy page dimensions from first section
    new_section.page_width = first_section.page_width
    new_section.page_height = first_section.page_height
    new_section.left_margin = first_section.left_margin
    new_section.right_margin = first_section.right_margin
    new_section.top_margin = first_section.top_margin
    new_section.bottom_margin = first_section.bottom_margin
    
    # Add "Page X of Y" footer to the new section
    _add_page_number_footer(new_section)
    _add_header_logo(new_section)


def _add_header_logo(section):
    """Add a small logo to the top-right of every page in this section."""
    header = section.header
    header.is_linked_to_previous = False
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_logo_if_present(para, width_in=0.95)


def add_table_of_contents(doc):
    """Add a Table of Contents page that auto-updates in Word."""
    heading = add_section_heading(doc, "Contents", font_size=18)

    # Add a TOC field — Word will populate page numbers when user presses F9
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    run2 = para.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-2" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = para.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar_separate)

    # Placeholder text (shown before TOC is updated in Word)
    run4 = para.add_run("Right-click and select 'Update Field' to populate contents.")
    run4.font.size = Pt(10)
    run4.font.italic = True
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run5 = para.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar_end)


def add_response_summary(doc, data):
    """Add response summary table."""
    heading = add_section_heading(doc, "Response Summary", font_size=14)

    response_counts = data.get('response_counts', {})

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set consistent widths (total ~6.1 inches to match PAPU-NANU tables)
    widths = content_columns(4.9, 1.1)
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Respondent Group"
    hdr_cells[0].width = widths[0]
    hdr_cells[1].text = "Responses"
    hdr_cells[1].width = widths[1]
    
    for cell in hdr_cells:
        set_cell_shading(cell, '183319')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
    
    for group in ['Self', 'Boss', 'Peers', 'DRs', 'Others']:
        if group in response_counts and response_counts[group] > 0:
            row = table.add_row().cells
            row[0].text = GROUP_DISPLAY[group]
            row[0].width = widths[0]
            row[1].text = str(response_counts[group])
            row[1].width = widths[1]
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    total = sum(response_counts.values())
    row = table.add_row().cells
    row[0].text = "Total"
    row[0].width = widths[0]
    row[0].paragraphs[0].runs[0].bold = True
    row[1].text = str(total)
    row[1].width = widths[1]
    row[1].paragraphs[0].runs[0].bold = True
    row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_fold_transparency_note(doc, data):
    """Generic disclosure that group-folding occurred, with no counts or group
    names. Fires for any tier of the cascade in database.py's
    get_leader_feedback_data (tier 1: Peers/DRs into Others; tier 2: Others
    into Peers/DRs; suppression, dormant fallback) via data['anonymity_applied'].
    This is now the SOLE fold-disclosure text in the report — the older
    per-group notes that used to live in add_response_summary (naming which
    groups were combined, or how many responses were suppressed) were removed
    2026-08-08 because keeping both alongside this generic one would name
    exactly what the generic wording is meant to hide. Silent when
    data['anonymity_applied'] is False, i.e. nothing folded.
    """
    if not data.get('anonymity_applied'):
        return

    note = doc.add_paragraph()
    run = note.add_run(
        "To protect anonymity, one or more respondent groups with too few "
        "responses to report individually have been combined with another "
        "group. Where this applies, the affected group label in this report "
        "reflects combined responses rather than a single respondent category."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_scoring_scale_note(doc, no_opportunity_label):
    """Explains what the 1-5 scores in this report actually measure. Placed
    once in About This Report, before any score appears - this is a linear,
    front-to-back read once per coaching conversation, not a navigable app,
    so one clear statement up front is enough; it isn't repeated near every
    chart or table.

    no_opportunity_label must be whichever "No opportunity to ..." wording is
    already correct for this report type (see feedback_form.py's self/other
    split): "No opportunity to demonstrate" for Self-Assessment, "No
    opportunity to observe" for Full 360. Don't invent a third variant.

    Hardcoded English for now, consistent with the rest of the report's
    content - this string will need the same _t()/get_translation()
    treatment as everything else once the i18n work resumes.
    """
    note = doc.add_paragraph()
    run = note.add_run(
        "Scores in this report reflect how often a behaviour is observed, "
        "not how well it's performed. 1 = Rarely or never · "
        "2 = Occasionally · 3 = Sometimes · 4 = Often · "
        f"5 = Consistently. Responses of '{no_opportunity_label}' are "
        "excluded from averages rather than counted as zero."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_executive_summary(doc, data):
    """Add executive summary with dimension table and radar chart."""
    add_section_heading(doc, "Executive Summary", font_size=16)
    
    # Dimension table with proper column widths
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Set column widths to match other tables (total ~6.1 inches)
    widths = content_columns(3.8, 0.7, 1.0, 0.5)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    
    hdr = table.rows[0].cells
    hdr[0].text = "Dimension"
    hdr[1].text = "Self"
    hdr[2].text = GROUP_DISPLAY['Combined']
    hdr[3].text = "Gap"

    for cell in hdr:
        set_cell_shading(cell, '183319')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    for dim_name in DIMENSIONS.keys():
        dim_data = data['by_dimension'].get(dim_name, {})
        row = table.add_row().cells
        
        # Set widths for each new row
        for i, cell in enumerate(row):
            cell.width = widths[i]
        
        row[0].text = dim_name
        row[1].text = f"{dim_data.get('Self', 0):.1f}" if dim_data.get('Self') else "-"
        row[2].text = f"{dim_data.get('Combined', 0):.1f}" if dim_data.get('Combined') else "-"
        
        # Centre-align the score columns
        for i in [1, 2, 3]:
            row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        gap = dim_data.get('Gap')
        if gap is not None:
            row[3].text = f"{gap:+.1f}"
            if gap > SIGNIFICANT_GAP:
                # Over-rating (self rates higher than others). Direction
                # confirmed against by_dimension['Gap'] = Self - Combined in
                # database.py (2026-08-08) - positive gap genuinely means
                # self > others here, not inverted. Deeper Heritage White
                # (COLOURS['gap_over_rating']), not the base heritage_white -
                # this table is scanned quickly in a live coaching
                # conversation, so more presence than the pale version wins.
                set_cell_shading(row[3], 'BBB8A3')  # gap_over_rating
            elif gap < -SIGNIFICANT_GAP:
                # Under-rating (others rate higher than self). 70% tint of
                # bentley_green (COLOURS['gap_under_rating']), same reasoning
                # as above - more saturated than the original pale pass.
                set_cell_shading(row[3], 'BAC2BA')  # gap_under_rating
        else:
            row[3].text = "-"

    # First occurrence of "All Raters" in the report (the same header appears
    # again in the radar legend just below, then throughout the Strengths &
    # Development tables and every item bar chart) - explain the composition
    # here once, rather than repeating it at every later occurrence.
    caption = doc.add_paragraph()
    run = caption.add_run(
        "All raters combined: line manager, peers, direct reports, and others."
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
        self_scores = {dim: data['by_dimension'].get(dim, {}).get('Self') for dim in DIMENSIONS}
        combined_scores = {dim: data['by_dimension'].get(dim, {}).get('Combined') for dim in DIMENSIONS}
        create_radar_chart(DIMENSIONS, self_scores, combined_scores, tmp.name)

        # Full-width radar on its own page. At CONTENT_WIDTH_IN the chart is
        # roughly 5in tall, and this page already carries the Response Summary and
        # the Executive Summary table, so it cannot share. Breaking deliberately
        # avoids Word leaving a ragged half-page gap and reads better anyway:
        # the numbers first, then the picture.
        chart_heading = doc.add_paragraph()
        chart_heading.paragraph_format.page_break_before = True
        chart_heading.paragraph_format.space_after = Pt(6)
        chart_heading.paragraph_format.keep_with_next = True
        run = chart_heading.add_run("Your Profile at a Glance")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(tmp.name, width=Inches(CONTENT_WIDTH_IN))
        os.unlink(tmp.name)

    # No explicit page break — next section uses page_break_before


def add_papu_nanu_section(doc, data):
    """Add strengths and development areas analysis."""
    heading = add_section_heading(doc, "Strengths & Development Analysis", font_size=16)
    heading.paragraph_format.page_break_before = True
    
    categories = categorize_papu_nanu(data)
    
    # Helper function to create consistent PAPU-NANU tables
    def keep_table_together(table):
        """Prevent table rows from splitting across pages."""
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.keep_with_next = True
        if table.rows:
            last_row = table.rows[-1]
            for cell in last_row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.keep_with_next = False
    
    def create_papu_table(doc, items, header_color, header_text_color=RGBColor(255, 255, 255)):
        """Create a PAPU-NANU table with consistent formatting.

        header_text_color defaults to white; Development Areas' heritage
        white background needs dark green text instead, since white text
        would be invisible on it.
        """
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False

        # Column widths: # | Behaviour | Self | All Raters | Gap
        widths = content_columns(0.4, 3.8, 0.6, 0.6, 0.6)

        hdr = table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "Behaviour"
        hdr[2].text = "Self"
        hdr[3].text = GROUP_DISPLAY['Combined']
        hdr[4].text = "Gap"

        for i, cell in enumerate(hdr):
            cell.width = widths[i]
            set_cell_shading(cell, header_color)
            cell.paragraphs[0].runs[0].font.color.rgb = header_text_color
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if i != 1:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for item in items:
            row = table.add_row().cells
            for i, cell in enumerate(row):
                cell.width = widths[i]
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after = Pt(2)
            
            row[0].text = str(item['item_num'])
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[0].paragraphs[0].runs[0].font.size = Pt(9)
            
            row[1].text = item['text']
            row[1].paragraphs[0].runs[0].font.size = Pt(9)
            
            row[2].text = f"{item['self']:.1f}"
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].paragraphs[0].runs[0].font.size = Pt(9)
            
            row[3].text = f"{item['combined']:.1f}"
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[3].paragraphs[0].runs[0].font.size = Pt(9)
            
            gap = item['gap']
            row[4].text = f"{gap:+.1f}" if gap is not None else "-"
            row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[4].paragraphs[0].runs[0].font.size = Pt(9)
        
        keep_table_together(table)
        return table
    
    # Agreed Strengths (Bentley green header)
    if categories['agreed_strengths']:
        heading = add_section_heading(doc, "Agreed Strengths", font_size=13, level=2)
        heading.paragraph_format.keep_with_next = True
        desc = doc.add_paragraph("You and others agree these are strengths - keep doing these.")
        desc.paragraph_format.keep_with_next = True
        create_papu_table(doc, categories['agreed_strengths'][:8], '183319')
        doc.add_paragraph()

    # Good News (a genuine tint of Bentley Green, not just "some other green" -
    # see 'good_news_tint' in framework.COLOURS)
    if categories['good_news']:
        heading = add_section_heading(doc, "Good News", font_size=13, level=2)
        heading.paragraph_format.keep_with_next = True
        desc = doc.add_paragraph("Others rate you higher than you rate yourself - you may be underselling yourself.")
        desc.paragraph_format.keep_with_next = True
        create_papu_table(doc, categories['good_news'], '5D705E')
        doc.add_paragraph()

    # Development Areas (heritage white header, dark green text - white text
    # would be invisible on this background)
    if categories['development_areas']:
        heading = add_section_heading(doc, "Development Areas", font_size=13, level=2)
        heading.paragraph_format.keep_with_next = True
        desc = doc.add_paragraph("Both you and others see room for growth - priority focus for development.")
        desc.paragraph_format.keep_with_next = True
        create_papu_table(
            doc, categories['development_areas'][:8], 'DCD8C0',
            header_text_color=RGBColor(0x18, 0x33, 0x19)
        )
        doc.add_paragraph()

    # Potential Blind Spots / Hidden Talents (charcoal grey header)
    if categories['hidden_talents']:
        heading = add_section_heading(doc, "Potential Blind Spots", font_size=13, level=2)
        heading.paragraph_format.keep_with_next = True
        desc = doc.add_paragraph(
            "You rate yourself higher than others do here. This might mean the strength "
            "isn't landing as clearly as you think, or that it's an area worth a closer look."
        )
        desc.paragraph_format.keep_with_next = True
        create_papu_table(doc, categories['hidden_talents'], '5F5E5A')
        doc.add_paragraph()
    
    # No explicit page break — next section uses page_break_before


def add_dimension_section(doc, dim_name, data, comments, is_self_only=False, is_first_dimension=False):
    """Add a dimension section with items displayed side-by-side with bar charts."""
    heading = add_section_heading(doc, dim_name, font_size=13, level=2)
    # Word's Heading 2 style carries its own 10pt space_before, which is
    # redundant here since this heading always starts at the very top of a
    # fresh page anyway (page_break_before below) - the page's own top
    # margin already provides that buffer.
    heading.paragraph_format.space_before = Pt(0)
    # Every dimension always starts on a fresh page (except the first, which
    # flows straight after the "Detailed Feedback" heading). A space-aware
    # heuristic was tried instead - only forcing a break when little room was
    # left on the current page - but comment length varies enough between
    # leaders that it produced worse results in practice (e.g. a dimension
    # starting at the very bottom of a page), so this reverts to the simple,
    # reliable rule per the human's instruction 2026-08-06.
    if not is_first_dimension:
        heading.paragraph_format.page_break_before = True

    # Dimension description. Tight space_after instead of the Normal default
    # (10pt) plus a whole extra blank paragraph after it (~23pt combined) -
    # one-time saving per dimension, small next to the per-item savings below
    # but free, since nothing readable shrinks. line_spacing=1.0 overrides
    # the document default (1.15 "auto"), shaving a proportional amount off
    # every wrapped line rather than just the paragraph's before/after -
    # matters most here since some dimension descriptions run to 3-4 lines.
    desc = doc.add_paragraph()
    desc.paragraph_format.space_before = Pt(0)
    desc.paragraph_format.space_after = Pt(6)
    desc.paragraph_format.line_spacing = 1.0
    run = desc.add_run(DIMENSION_DESCRIPTIONS[dim_name])
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    start, end = DIMENSIONS[dim_name]
    
    # Each item: side-by-side borderless table (text left, bar chart right)
    for item_num in range(start, end + 1):
        item_scores = data['by_item'].get(item_num, {})

        # Serve the wording the reader actually answered. In the Self-Assessment
        # report the only rater was the leader, so the I-form is correct; in the
        # Full 360 the same item was put to others as the They-form, so showing
        # the I-form there would misrepresent what a peer or direct report was
        # asked. NB: do NOT fall back to item_scores['text'] here, because
        # get_leader_feedback_data bakes the They-form into it unconditionally.
        item_text = get_item_text(item_num, 'Self' if is_self_only else 'Others')

        layout_table = doc.add_table(rows=1, cols=2)
        make_table_borderless(layout_table)
        layout_table.autofit = False

        # Chart column is the fixed ITEM_CHART_WIDTH_IN (see its definition -
        # deliberately not proportional, so chart height stays constant
        # regardless of page geometry); text column takes whatever's left of
        # CONTENT_WIDTH_IN, so a wider page gives the text more room to
        # wrap into fewer lines rather than stretching the chart taller.
        item_col_widths = [
            Inches(CONTENT_WIDTH_IN - ITEM_CHART_WIDTH_IN),
            Inches(ITEM_CHART_WIDTH_IN),
        ]

        # Prevent a single item's row (text + chart) from ever splitting
        # across a page boundary - a defensive backstop, not the primary
        # mechanism (that's getting the sizing right below), but real Word
        # text-wrapping can still vary slightly from what's calculated here.
        tr = layout_table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)

        # Every paragraph in this block gets an explicit, tight space_before/
        # after instead of inheriting Normal's defaults (0pt before, 10pt
        # after) - fitting five items per dimension page is a spacing
        # problem, not a font-size one (see the 2026-08-06 font audit: item
        # text and comment body are already tied for smallest visible text
        # in the document bar the 8/9pt captions, so there's little headroom
        # left to shrink text further without hurting legibility).
        text_cell = layout_table.rows[0].cells[0]
        text_cell.width = item_col_widths[0]
        text_para = text_cell.paragraphs[0]
        text_para.add_run(f"Q{item_num}. ").bold = True
        text_para.add_run(item_text)
        text_para.runs[0].font.size = Pt(10)
        if len(text_para.runs) > 1:
            text_para.runs[1].font.size = Pt(10)
        text_para.paragraph_format.space_before = Pt(0)
        text_para.paragraph_format.space_after = Pt(2)
        text_para.paragraph_format.line_spacing = 1.0
        text_para.paragraph_format.keep_with_next = True

        # Whole-item coverage — never per-group (see anonymity design principle)
        if not is_self_only:
            total_respondents = sum(data.get('raw_response_counts', {}).values())
            no_opp_count = data.get('no_opportunity', {}).get(item_num, {}).get('count', 0)
            rated = total_respondents - no_opp_count
            coverage_para = text_cell.add_paragraph()
            coverage_run = coverage_para.add_run(f"Rated by {rated} of {total_respondents} respondents")
            coverage_run.font.size = Pt(8)
            coverage_run.font.italic = True
            coverage_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            coverage_para.paragraph_format.space_before = Pt(0)
            coverage_para.paragraph_format.space_after = Pt(2)
            coverage_para.paragraph_format.line_spacing = 1.0
            coverage_para.paragraph_format.keep_with_next = True

        chart_cell = layout_table.rows[0].cells[1]
        chart_cell.width = item_col_widths[1]

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            if is_self_only:
                create_self_only_bar(item_scores.get('Self'), tmp.name)
            else:
                create_item_bar_chart(item_scores, tmp.name)

            chart_para = chart_cell.paragraphs[0]
            chart_para.add_run().add_picture(tmp.name, width=Inches(ITEM_CHART_WIDTH_IN))
            chart_para.paragraph_format.space_before = Pt(0)
            # 0 not 2: the spacer paragraph immediately after this cell's
            # table carries the whole gap to the next item on its own now
            # (see below) - having both add space_after was double-counting
            # the same visual gap.
            chart_para.paragraph_format.space_after = Pt(0)
            chart_para.paragraph_format.keep_together = True
            os.unlink(tmp.name)

        # Spacer between this item's table and the next - tight space_after
        # rather than the Normal default (10pt plus the paragraph's own line
        # height), so five of these per dimension don't add up to real
        # wasted space. Still a real (non-empty-looking) paragraph, not
        # removed outright, so item boundaries stay unambiguous at the XML
        # level the way coverage_para etc. already rely on.
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(3)
        spacer.paragraph_format.line_spacing = 1.0

    # --- CLEAN COMMENTS (replaces old table style) ---
    # The Self-Assessment report must only ever show the leader's own comment,
    # never other raters' - it's generated before Module 1, when the leader is
    # the only person who has responded at all, and even when regenerated later
    # for the record it should still read as the leader's own view only, not a
    # mix of the two report stages.
    section_comments = comments.get('by_section', {}).get(dim_name, [])
    if is_self_only:
        section_comments = [c for c in section_comments if c.get('group') == 'Self']
    if section_comments:
        comment_heading = doc.add_paragraph()
        # Deliberate, unconditional break for the FULL 360 case - comments
        # always start on a fresh page there, whether there's 2 short comments
        # or 20 long ones from five different rater groups. Structural
        # consistency (the same element always behaving the same way) is the
        # goal, not using up whatever space happens to be left after the five
        # items above - that's what made the previous space-aware heuristic
        # for dimension breaks unpredictable, and this deliberately avoids
        # repeating that mistake in a new spot.
        #
        # The Self-Assessment case is different enough to warrant a different
        # rule: there is ALWAYS at most one comment (the leader's own; see the
        # filter above), never an unbounded mix from several groups, so the
        # volume this has to absorb is bounded and small. Forcing a fresh page
        # for one short paragraph wastes it outright. No break here lets Word
        # place it on the same page as the five items whenever it genuinely
        # fits - not an estimate, just the natural flow - and only spill to a
        # new page when it doesn't.
        if not is_self_only:
            comment_heading.paragraph_format.page_break_before = True
        comment_heading.paragraph_format.keep_with_next = True
        run = comment_heading.add_run(f"Comments on {dim_name}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)

        add_clean_comments(doc, section_comments)

    # No explicit page break here — next section uses page_break_before


def add_overall_comments(doc, comments):
    """Add overall qualitative feedback section — clean style."""
    heading = add_section_heading(doc, "Overall Qualitative Feedback", font_size=16)
    heading.paragraph_format.page_break_before = True

    # --- KEEP DOING ---
    if comments.get('keep'):
        heading = doc.add_heading("What to Keep Doing", level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)

        add_clean_comments(doc, comments['keep'])

    doc.add_paragraph()

    # --- CHANGE ---
    if comments.get('change'):
        heading = doc.add_heading("The One Change That Would Make the Biggest Difference", level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)

        add_clean_comments(doc, comments['change'])


def add_development_priorities(doc, priorities, data, is_self_only=False):
    """
    Add the leader's self-identified development priorities.

    Captured at self-assessment, before any feedback arrived. In the Full 360
    the stated priority is shown alongside how that dimension actually scored,
    which is the point: it shows where the leader's own intent and the feedback
    agree, and where they diverge.

    The section renders even when nothing was stored (a leader who submitted
    before at least one priority became compulsory), because the capture space
    for coaching additions is worth having either way.
    """
    stated = [p for p in priorities if p.get('dimension')]

    heading = add_section_heading(doc, "Your Development Priorities", font_size=16)
    heading.paragraph_format.page_break_before = True

    if not stated:
        intro_text = (
            "You did not record development priorities with your self-assessment. "
            "Use the space below to set them out with your coach."
        )
    elif is_self_only:
        intro_text = (
            "You named these priorities when you completed your self-assessment, "
            "before receiving any feedback."
        )
    else:
        intro_text = (
            "You named these priorities when you completed your self-assessment, "
            "before receiving any feedback. The scores alongside each one show how "
            "that dimension was actually rated, so you can see where your own view "
            "and your feedback point the same way."
        )

    intro = doc.add_paragraph()
    run = intro.add_run(intro_text)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for priority in stated:
        dimension = priority['dimension']

        # Priority heading with rank
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(10)
        title_para.paragraph_format.space_after = Pt(2)
        title_para.paragraph_format.keep_with_next = True
        run = title_para.add_run(f"Priority {priority.get('rank', '')}: {dimension}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)

        # How that dimension actually scored
        dim_data = data.get('by_dimension', {}).get(dimension, {})
        self_score = dim_data.get('Self')
        combined = dim_data.get('Combined')
        gap = dim_data.get('Gap')

        score_bits = []
        if self_score is not None:
            score_bits.append(f"Your rating: {self_score:.1f}")
        if not is_self_only and combined is not None:
            score_bits.append(f"Others: {combined:.1f}")
            if gap is not None:
                score_bits.append(f"Gap: {gap:+.1f}")

        if score_bits:
            score_para = doc.add_paragraph()
            score_para.paragraph_format.space_before = Pt(0)
            score_para.paragraph_format.space_after = Pt(4)
            score_para.paragraph_format.keep_with_next = True
            run = score_para.add_run("   ·   ".join(score_bits))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x4A, 0x55, 0x68)

        # What the leader said they would work on
        actions = priority.get('actions')
        if actions:
            action_para = doc.add_paragraph()
            action_para.paragraph_format.space_before = Pt(0)
            action_para.paragraph_format.space_after = Pt(8)
            run = action_para.add_run(actions)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

        _add_thin_rule(doc)

    # --- Space to ADD priorities during the coaching conversation ---
    # Additive by design: the coaching conversation builds on what the leader
    # named rather than replacing it, and the stored set stays as the
    # pre-feedback baseline the Full 360 report compares against.
    doc.add_paragraph()

    add_heading = doc.add_paragraph()
    add_heading.paragraph_format.space_before = Pt(10)
    add_heading.paragraph_format.space_after = Pt(2)
    add_heading.paragraph_format.keep_with_next = True
    run = add_heading.add_run(
        "Your Priorities" if not stated else "Adding to Your Priorities"
    )
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)

    if not stated:
        prompt_text = (
            "Use this space in your coaching conversation to set out the areas you "
            "want to work on and what you will do differently."
        )
    elif is_self_only:
        prompt_text = (
            "Use this space in your coaching conversation to add anything further you "
            "want to work on. These add to the priorities above rather than replacing "
            "them, so the record of what you set out to change stays intact."
        )
    else:
        prompt_text = (
            "Use this space to add anything your feedback has surfaced that you now "
            "want to work on. These add to the priorities above rather than replacing "
            "them, so you keep sight of what you set out to change before the feedback "
            "arrived."
        )

    add_writing_prompt(doc, prompt_text)
    add_priority_capture_table(doc, rows=3)


def add_reflection_questions(doc):
    """
    Add reflection questions with space to answer them.

    Each question gets ruled lines: a prompt with nowhere to write is a wasted
    prompt, and this report is meant to be worked on rather than filed.
    """
    heading = doc.add_heading("Reflection Questions", level=1)
    heading.paragraph_format.page_break_before = True

    doc.add_paragraph(
        "Use these to prepare for your coaching conversation, and to capture what comes out of it. "
        "Write straight onto this document."
    )

    questions = [
        "Which dimensions did you rate yourself highest on? What evidence supports these ratings?",
        "Which dimensions did you rate yourself lowest on? What makes these areas challenging?",
        "Were there any items where you found it difficult to decide on a rating? What made them difficult?",
        "Looking at the priorities you named, why do those matter most to you right now?",
        "What support or resources might help you develop in these areas?",
        "Where do you expect others' views of your leadership to differ from your own?",
    ]

    for i, question in enumerate(questions, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.keep_with_next = True
        para.add_run(f"{i}. ").bold = True
        para.add_run(question)
        add_writing_lines(doc, count=3)


def add_what_happens_next(doc):
    """
    Add What Happens Next section for self-assessment reports.

    The sequence matters and must match the real programme timeline: the leader
    reads this at Module 1, having just received the report and BEFORE they have
    nominated anyone. An earlier version described feedback collection as already
    under way and placed coaching after the full report, which would have told
    the reader they had missed a step.
    """
    heading = doc.add_heading("What Happens Next", level=1)
    heading.paragraph_format.page_break_before = True

    doc.add_paragraph(
        "This self-assessment is the first of two stages. Here is how the rest "
        "of the process runs:"
    )

    steps = [
        ("Your Coaching Conversation (following Module 1)",
         "You will talk this report through with your coach. The aim is to understand your own view "
         "of your leadership first, before any other perspectives are introduced, and to sharpen the "
         "development priorities you have named."),
        ("Nominating Your Respondents (between Modules 1 and 2)",
         "You will be invited to nominate the people who will give you feedback: your line manager, "
         "peers, direct reports, and any others who see you lead regularly. Your coaching conversation "
         "will help you think about who will give you the most useful view."),
        ("Feedback Collection",
         "Those you nominate receive their own confidential link. You will be able to see overall "
         "progress in your portal, though not who has or has not responded."),
        ("Your Full Feedback Report (Module 2)",
         "Once enough responses are in, your full 360 report is produced. It sets others' perceptions "
         "alongside the self-assessment in this document, which is where the areas of agreement and "
         "the blind spots become visible."),
        ("Development Planning",
         "Working with your coach, you will build a focused plan from what the feedback shows, adding "
         "to the priorities you have already set rather than starting again."),
    ]

    for i, (title, description) in enumerate(steps, 1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        para.add_run(f"{i}. {title}: ").bold = True
        para.add_run(description)
        doc.add_paragraph()


def synthesise_feedback_themes(leader_name, comments, data):
    """
    Use the Claude API to synthesise key themes from all verbatim feedback.

    Returns (themes, warning):
      themes: list of theme dicts [{'title': str, 'narrative': str}, ...], or None.
      warning: None on success or on the benign "too few comments" skip; a
        human-readable string when synthesis was attempted but failed, so the
        admin UI can surface it instead of the section silently vanishing.
    """
    import json
    
    # Collect all comments into a structured prompt
    all_comments = []
    
    # Dimension comments
    for dim_name, dim_comments in comments.get('by_section', {}).items():
        for c in dim_comments:
            all_comments.append({
                'dimension': dim_name,
                'source': c['group'],
                'text': c['text']
            })
    
    # Overall keep/change comments
    for c in comments.get('keep', []):
        all_comments.append({
            'dimension': 'What to Keep Doing',
            'source': c['group'],
            'text': c['text']
        })
    for c in comments.get('change', []):
        all_comments.append({
            'dimension': 'The One Change',
            'source': c['group'],
            'text': c['text']
        })
    
    # Need enough comments to synthesise meaningfully. Not a failure, just not
    # enough material yet, so no warning.
    if len(all_comments) < 5:
        return None, None
    
    # Build dimension scores context. Self/Combined/Gap is the self-vs-others
    # picture; the per-group breakdown (Boss/Peers/DRs/Others) additionally
    # surfaces whether the OTHER raters agree with each other, which a single
    # Combined figure averages away. Only included when 2+ groups are
    # individually visible AND they actually diverge (spread >= 0.5) - a tight
    # cluster isn't a theme, and forcing one in would read as manufactured.
    # These per-group averages are exactly what's already in data['by_dimension']
    # from get_leader_feedback_data's own anonymity fold-cascade (a group only
    # has its own key there if it already cleared the threshold, or was never
    # subject to it like Boss), so this adds no new disclosure - it's the same
    # numbers already shown elsewhere in the report, just previously withheld
    # from this one section.
    scores_context = []
    for dim_name in DIMENSIONS.keys():
        dim_data = data.get('by_dimension', {}).get(dim_name, {})
        self_score = dim_data.get('Self')
        combined = dim_data.get('Combined')
        gap = dim_data.get('Gap')
        if not combined:
            continue

        # No self-score (leader hasn't self-assessed yet) gets its own
        # cleaner format rather than a literal "Self=None" in the prompt.
        if self_score is not None:
            line = f"{dim_name}: Self={self_score}, Others={combined}"
            if gap:
                line += f", Gap={gap:+.1f}"
        else:
            line = f"{dim_name}: Combined={combined}"

        group_scores = {g: dim_data[g] for g in ('Boss', 'Peers', 'DRs', 'Others') if g in dim_data}
        if len(group_scores) >= 2:
            spread = max(group_scores.values()) - min(group_scores.values())
            if spread >= 0.5:
                by_group = ", ".join(f"{g}={v}" for g, v in group_scores.items())
                line += f" | By group: {by_group} (spread={spread:.1f})"

        scores_context.append(line)

    prompt = f"""You are writing a section of a 360-degree feedback report for {leader_name}. This section synthesises the key themes from their verbatim feedback. The leader will read this directly, so write in the second person — use "you", "your", "your feedback suggests" etc. Never refer to them as "the leader" or "this leader".

The tone should be warm, constructive, and developmental — as if a skilled coach were talking them through their feedback. Be direct but supportive.

Below are all the verbatim comments from their feedback, organised by dimension and source group, followed by their dimension scores.

VERBATIM COMMENTS:
{json.dumps(all_comments, indent=2)}

DIMENSION SCORES:
{chr(10).join(scores_context)}

Please identify 4-6 key themes that emerge from this feedback. For each theme:
1. Give it a clear, concise title (e.g., "Building Trust Through Authenticity" or "Balancing Operational Focus with Strategic Thinking")
2. Write a 2-3 sentence narrative that synthesises the evidence — reference what respondents said without quoting them verbatim, connect to the quantitative scores where relevant, and speak directly to {leader_name} using "you" and "your".

Where a dimension shows scores by individual rater group, consider whether those groups diverge meaningfully from each other, not only from the self-rating. Genuine disagreement between groups (e.g. Direct Reports and Peers seeing something differently) is a distinct, often more actionable insight than a self-vs-others gap, and is worth its own theme if the spread is notable. Don't manufacture a divergence theme where the groups broadly agree."""

    request_payload = {
        "model": SYNTHESIS_MODEL,
        # Thinking is ON BY DEFAULT on this model, and max_tokens caps
        # thinking PLUS response text together. 2000 was enough when the
        # old model did no thinking; it would now truncate the JSON.
        "max_tokens": 8000,
        # Structured outputs guarantee valid JSON, which removes the need
        # to strip markdown fences off the response and hope it parses.
        "output_config": {
            "format": {
                "type": "json_schema",
                "schema": {
                    "type": "object",
                    "properties": {
                        "themes": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "title": {"type": "string"},
                                    "narrative": {"type": "string"},
                                },
                                "required": ["title", "narrative"],
                                "additionalProperties": False,
                            },
                        }
                    },
                    "required": ["themes"],
                    "additionalProperties": False,
                },
            }
        },
        "messages": [
            {"role": "user", "content": prompt}
        ]
    }

    for attempt in range(SYNTHESIS_MAX_RETRIES + 1):
        try:
            response = requests.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "Content-Type": "application/json",
                    "x-api-key": _get_api_key(),
                    "anthropic-version": "2023-06-01"
                },
                json=request_payload,
                timeout=120
            )

            if response.status_code == 200:
                result = response.json()

                # A refusal returns HTTP 200 with no usable content, so check
                # before reading content[0]
                if result.get('stop_reason') == 'refusal':
                    warning = "the request was declined by the Anthropic API's safety classifiers"
                    print(f"SYNTHESIS SKIPPED: {warning}. The rest of the report is unaffected.",
                          file=sys.stderr)
                    return None, warning

                # content[0] is NOT reliably the text block: thinking is on by
                # default for this model (see max_tokens note above), and a
                # thinking response puts a {"type": "thinking", ...} block
                # first in content, with no "text" key at all - indexing [0]
                # directly raised KeyError: 'text' in production the first
                # time this path actually ran against the live API (never
                # caught locally, since this session had no API key to
                # exercise a real call). Find the actual text block instead
                # of assuming its position.
                text_block = next(
                    (block for block in result.get('content', []) if block.get('type') == 'text'),
                    None
                )
                if text_block is None:
                    warning = "the API response contained no text content (thinking only)"
                    print(f"SYNTHESIS FAILED: {warning}.", file=sys.stderr)
                    return None, warning

                text = text_block['text'].strip()
                return json.loads(text)['themes'], None

            # 429/500/529 are transient and routinely resolve a few seconds
            # later, so worth a couple of quick retries rather than losing the
            # whole section over a momentary blip. Anything else (bad key, bad
            # request) won't be fixed by retrying, so fails immediately.
            if response.status_code in SYNTHESIS_TRANSIENT_STATUS_CODES and attempt < SYNTHESIS_MAX_RETRIES:
                print(f"SYNTHESIS RETRY {attempt + 1}/{SYNTHESIS_MAX_RETRIES}: HTTP "
                      f"{response.status_code} from the Anthropic API (transient), "
                      f"retrying in {SYNTHESIS_RETRY_DELAY_SECONDS}s...", file=sys.stderr)
                time.sleep(SYNTHESIS_RETRY_DELAY_SECONDS)
                continue

            # Loud, specific, and on stderr so it reaches the Streamlit logs,
            # and returned as a warning so the admin UI can surface it too. A
            # silent skip here means a leader's report quietly loses a whole
            # section, which is worse than a visible failure.
            warning = f"the Anthropic API returned HTTP {response.status_code}"
            print(f"SYNTHESIS FAILED: {warning}. The Key Themes section will be "
                  f"missing from this report. Response: {response.text[:500]}", file=sys.stderr)
            return None, warning

        except Exception as e:
            print(f"Theme synthesis failed: {e}")
            return None, f"theme synthesis raised an exception: {e}"


def _get_api_key():
    """
    Get the Anthropic API key from the environment, falling back to
    Streamlit secrets.

    Environment variable checked first (2026-08-06, was the other way round)
    so the app isn't tied to Streamlit Cloud's specific secrets mechanism -
    we've moved hosting once already and might reasonably do so again. A
    plain ANTHROPIC_API_KEY env var works the same way on Render, Streamlit
    Cloud, or anywhere else; the Streamlit secrets path only matters for a
    deployment that still configures it that way instead.
    """
    key = os.environ.get("ANTHROPIC_API_KEY")
    if key:
        return key

    # Fall back to Streamlit secrets, structured as [anthropic] api_key = "..."
    # (not a flat ANTHROPIC_API_KEY key) - matches how existing Streamlit
    # Cloud deployments of this app already have their secrets.toml set up.
    try:
        import streamlit as st
        key = st.secrets.get("anthropic", {}).get("api_key")
        if key:
            return key
    except Exception:
        pass

    raise ValueError("No Anthropic API key found. Set ANTHROPIC_API_KEY or add to Streamlit secrets.")


def add_theme_synthesis(doc, leader_name, comments, data):
    """
    Add the AI-generated theme synthesis section to the report.

    Falls back gracefully if the API is unavailable — the report still generates
    without the synthesis section. Returns a warning string when the section was
    skipped due to a genuine failure (as opposed to too few comments to
    synthesise), so the caller can surface it instead of the section silently
    vanishing.
    """
    themes, warning = synthesise_feedback_themes(leader_name, comments, data)

    if not themes:
        return warning
    
    heading = add_section_heading(doc, "Key Themes in Your Feedback", font_size=16)
    heading.paragraph_format.page_break_before = True
    
    intro = doc.add_paragraph(
        "The following themes have been identified from the qualitative feedback provided by your "
        "respondents. These represent the patterns and consistent messages that emerge when all "
        "comments are considered together."
    )
    intro.paragraph_format.space_after = Pt(12)
    
    for i, theme in enumerate(themes):
        # Theme title
        title_para = doc.add_paragraph()
        title_para.paragraph_format.space_before = Pt(12) if i > 0 else Pt(6)
        title_para.paragraph_format.space_after = Pt(4)
        run = title_para.add_run(f"{i + 1}. {theme['title']}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x18, 0x33, 0x19)  # Bentley green
        
        # Theme narrative
        narrative_para = doc.add_paragraph(theme['narrative'])
        narrative_para.paragraph_format.space_after = Pt(8)
        for run in narrative_para.runs:
            run.font.size = Pt(10)
    
    # Closing note
    _add_thin_rule(doc)
    note = doc.add_paragraph()
    run = note.add_run(
        "Note: This synthesis has been generated to help identify patterns in your feedback. "
        "It should be explored in your coaching session alongside the detailed verbatim comments above."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    return None


def add_next_steps(doc):
    """Add next steps section for full 360 reports."""
    heading = add_section_heading(doc, "Next Steps", font_size=16)
    heading.paragraph_format.page_break_before = True
    
    doc.add_paragraph(
        "This feedback provides a foundation for your ongoing leadership development. "
        "Consider the following as you reflect on your results:"
    )
    
    steps = [
        "Review your Agreed Strengths - how can you leverage these more deliberately?",
        "Consider the Good News items - are you being too hard on yourself in these areas?",
        "Focus on 2-3 Development Areas - what specific actions could you take?",
        "Explore any Potential Blind Spots - is this a visibility issue or a genuine gap?",
        "Discuss your results with your coach to create a focused development plan.",
    ]
    
    for step in steps:
        para = doc.add_paragraph(step, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        "Remember: this feedback represents perceptions at a point in time. "
        "Use it as data to inform your development, not as a definitive judgement."
    )

    # Space to work in, so this report is a live document rather than a read-only output
    add_writing_prompt(
        doc,
        "What stands out most from this feedback, and what will you do about it? "
        "Write your thinking here before your coaching conversation."
    )
    add_writing_lines(doc, count=8)


# ============================================
# MAIN GENERATION
# ============================================

def generate_report(leader_name, report_type, data, comments, dealership=None, cohort=None):
    """
    Main entry point for report generation.
    
    Args:
        leader_name: Name of the leader
        report_type: 'Self-Assessment', 'Full 360', or 'Progress Report'
        data: Feedback data dictionary
        comments: Comments dictionary
        dealership: Optional dealership name
        cohort: Optional cohort name
    
    Returns:
        (output_path, theme_warning) — theme_warning is None unless the Key
        Themes section (Full 360 only) was skipped due to a genuine failure.
    """
    doc = Document()
    apply_page_geometry(doc)
    apply_document_font(doc)
    theme_warning = None

    if report_type == 'Self-Assessment':
        create_cover_page(doc, leader_name, "Self-Assessment Report", dealership, cohort)

        # Contents page leads, with real page numbers once updated in Word —
        # About This Report follows and no longer needs to repeat the section
        # list, since Contents already covers that.
        add_table_of_contents(doc)

        about_heading = add_section_heading(doc, "About This Report", font_size=16)
        about_heading.paragraph_format.page_break_before = True
        doc.add_paragraph(
            "This report captures your own view of your leadership effectiveness across the nine "
            "dimensions of the Compass framework. It is the first of two stages: this "
            "self-assessment, which you talk through with your coach at Module 1, and your full "
            "360 feedback report at Module 2, once the people you nominate have given their view."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            "Write on this document. There is space throughout to capture what comes out of your "
            "coaching conversation, and it is yours to keep working on."
        )
        doc.add_paragraph()
        add_scoring_scale_note(doc, "No opportunity to demonstrate")

        # Overview — dimension table + radar on one page
        heading = add_section_heading(doc, "Your Self-Assessment Overview", font_size=16)
        heading.paragraph_format.page_break_before = True
        
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        
        widths = content_columns(4.5, 1.5)
        
        hdr = table.rows[0].cells
        hdr[0].text = "Dimension"
        hdr[1].text = "Your Score"
        for i, cell in enumerate(hdr):
            cell.width = widths[i]
            set_cell_shading(cell, '183319')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
        if len(hdr) > 1:
            hdr[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for dim_name in DIMENSIONS.keys():
            row = table.add_row().cells
            for i, cell in enumerate(row):
                cell.width = widths[i]
            row[0].text = dim_name
            self_score = data['by_dimension'].get(dim_name, {}).get('Self')
            row[1].text = f"{self_score:.1f}" if self_score else "-"
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Keep table with radar. This page carries only the heading and a ten-row
        # table, so a full-width radar (about 5in tall) still fits beneath it.
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.keep_with_next = True

        # Horizontal rule + radar chart
        _add_thin_rule(doc)

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            self_scores = {dim: data['by_dimension'].get(dim, {}).get('Self') for dim in DIMENSIONS}
            create_radar_chart(DIMENSIONS, self_scores, None, tmp.name)

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(tmp.name, width=Inches(CONTENT_WIDTH_IN))
            os.unlink(tmp.name)
        
        # Detailed sections — with parent heading that flows into first dimension
        detail_heading = add_section_heading(doc, "Detailed Self-Assessment by Dimension", font_size=18)
        detail_heading.paragraph_format.page_break_before = True
        detail_heading.paragraph_format.keep_with_next = True
        
        for i, dim_name in enumerate(DIMENSIONS.keys()):
            add_dimension_section(doc, dim_name, data, comments, is_self_only=True,
                                  is_first_dimension=(i == 0))

        # Self-identified development priorities
        add_development_priorities(
            doc, data.get('development_priorities', []), data, is_self_only=True
        )

        # Reflection Questions
        add_reflection_questions(doc)
        
        # What Happens Next
        add_what_happens_next(doc)
        
    elif report_type == 'Full 360':
        create_cover_page(doc, leader_name, "Feedback Report", dealership, cohort)

        # Contents page leads, with real page numbers once updated in Word —
        # About This Report follows and no longer needs to repeat the section
        # list, since Contents already covers that.
        add_table_of_contents(doc)

        about_heading = add_section_heading(doc, "About This Report", font_size=16)
        about_heading.paragraph_format.page_break_before = True
        doc.add_paragraph(
            "This 360-degree feedback report brings together perspectives from your line manager, "
            "peers, direct reports, and others, alongside your self-assessment. The comparison "
            "helps identify areas of alignment and potential blind spots."
        )
        doc.add_paragraph()
        add_scoring_scale_note(doc, "No opportunity to observe")

        # Response Summary + Executive Summary + Radar — all on one page
        add_response_summary(doc, data)
        add_fold_transparency_note(doc, data)
        add_executive_summary(doc, data)
        
        add_papu_nanu_section(doc, data)
        
        # Detailed Feedback by Dimension
        # The section title is added before the first dimension only
        detail_heading = add_section_heading(doc, "Detailed Feedback by Dimension", font_size=18)
        detail_heading.paragraph_format.page_break_before = True
        # keep_with_next ensures heading stays with first dimension
        detail_heading.paragraph_format.keep_with_next = True
        
        for i, dim_name in enumerate(DIMENSIONS.keys()):
            add_dimension_section(doc, dim_name, data, comments, is_self_only=False,
                                  is_first_dimension=(i == 0))

        add_overall_comments(doc, comments)
        theme_warning = add_theme_synthesis(doc, leader_name, comments, data)
        add_development_priorities(
            doc, data.get('development_priorities', []), data, is_self_only=False
        )
        add_next_steps(doc)
    
    else:  # Progress Report
        create_cover_page(doc, leader_name, "Progress Report", dealership, cohort)
        doc.add_paragraph("Progress report generation coming soon...")
    
    # Save. The old filename (name_type_YYYYMMDD.docx, date only, no time or
    # unique component) collided whenever the same leader+report type was
    # generated more than once on the same calendar day - trivially easy to
    # trigger with a double-click on "Generate" before the first request
    # finishes, since REPORTS_DIR is one shared folder on the server, not
    # per-session. Two concurrent doc.save() calls to the identical path can
    # corrupt each other mid-write, and admin_dashboard.py's subsequent
    # open(output_path, 'rb') can catch the file in that truncated state -
    # this is what produced a real download link pointing at a near-empty
    # file. A uuid suffix makes every generation's path unique regardless of
    # leader, type, or timing, so concurrent/repeated generations can never
    # collide.
    unique_id = uuid.uuid4().hex[:8]
    filename = f"{leader_name.replace(' ', '_')}_{report_type.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}_{unique_id}.docx"
    output_path = REPORTS_DIR / filename
    doc.save(output_path)

    return str(output_path), theme_warning


def generate_all_reports(db, leader_ids=None):
    """Generate reports for multiple leaders."""
    
    if leader_ids is None:
        leaders = db.get_all_leaders()
        leader_ids = [l['id'] for l in leaders if l['completed_raters'] >= 5]
    
    generated = []
    for leader_id in leader_ids:
        leader = db.get_leader(leader_id)
        data, comments = db.get_leader_feedback_data(leader_id)
        
        output_path, theme_warning = generate_report(
            leader['name'],
            'Full 360',
            data,
            comments,
            leader.get('dealership'),
            leader.get('cohort')
        )
        if theme_warning:
            print(f"Key Themes could not be generated for {leader['name']}: {theme_warning}", file=sys.stderr)
        generated.append(output_path)
    
    return generated
